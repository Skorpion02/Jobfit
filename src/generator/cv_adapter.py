from typing import Dict, List
import logging
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Importar LM Studio client
try:
    from src.llm.lmstudio_client import lmstudio_client
except ImportError:
    lmstudio_client = None


class CVAdapter:
    """
    Adaptador de CV que personaliza el CV según los requisitos de la oferta
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def adapt_cv(self, cv: Dict, job_offer: Dict, matching_results: Dict) -> Dict:
        """
        Adapta un CV para una oferta específica basándose en los resultados del matching
        """
        try:
            # Defensive normalization: asegurarnos de que los campos esperados
            # tengan tipos consistentes (listas/dicts) para evitar errores
            # 'NoneType' is not iterable durante la adaptación.
            if cv is None:
                cv = {}

            # Normalizar experience, education y projects a listas
            if not isinstance(cv.get('experience'), list):
                cv['experience'] = cv.get('experience') or []
            if not isinstance(cv.get('education'), list):
                cv['education'] = cv.get('education') or []
            if not isinstance(cv.get('projects'), list):
                cv['projects'] = cv.get('projects') or []

            # Normalizar skills a dict con claves 'technical' y 'other'
            skills_field = cv.get('skills')
            if isinstance(skills_field, list):
                cv['skills'] = {'technical': skills_field, 'other': []}
            elif isinstance(skills_field, dict):
                # Asegurar listas dentro del dict
                tech = skills_field.get('technical') or []
                other = skills_field.get('other') or []
                # Si alguien guardó por error un string, convertirlo a lista
                if isinstance(tech, str):
                    tech = [tech]
                if isinstance(other, str):
                    other = [other]
                cv['skills'] = {'technical': tech, 'other': other}
            else:
                cv['skills'] = {'technical': [], 'other': []}

            adapted_cv = {
                'personal_info': self._adapt_personal_info(cv.get('personal_info', {})),
                'summary': self._generate_targeted_summary(cv, job_offer, matching_results),
                'experience': self._prioritize_experience(cv.get('experience', []), matching_results),
                'skills': self._highlight_relevant_skills(cv.get('skills', {}), matching_results),
                'education': cv.get('education', []),
                'projects': self._highlight_relevant_projects(cv.get('projects', []), matching_results),
                'adaptation_notes': self._generate_adaptation_notes(matching_results)
            }
            
            self.logger.info(f"CV adaptado para: {job_offer.get('title', 'Posición sin título')}")
            return adapted_cv
            
        except Exception as e:
            self.logger.error(f"Error adaptando CV: {e}")
            return cv
    
    def export_to_docx(self, adapted_cv: Dict, filename: str = "cv_adaptado.docx") -> str:
        """Exporta el CV adaptado a formato DOCX con diseño profesional"""
        doc = Document()
        
        # Configurar márgenes del documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        personal_info = adapted_cv.get('personal_info', {})
        
        # === ENCABEZADO PROFESIONAL ===
        if personal_info.get('name'):
            # Nombre principal con formato elegante
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(personal_info['name'].upper())
            name_run.font.size = Pt(18)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(31, 78, 121)  # Azul profesional
            
            # Título profesional si existe
            title = personal_info.get('title', 'Profesional Especializado')
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(12)
            title_run.font.italic = True
            title_run.font.color.rgb = RGBColor(89, 89, 89)
            
            # Línea separadora
            self._add_horizontal_line(doc)
        
        # Información de contacto en formato profesional
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(f"Email: {personal_info['email']}")
        if personal_info.get('phone'):
            contact_info.append(f"Teléfono: {personal_info['phone']}")
        if personal_info.get('location'):
            contact_info.append(f"Ubicación: {personal_info['location']}")
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(89, 89, 89)
            
            # Espacio después del encabezado
            doc.add_paragraph()
        
        # === RESUMEN PROFESIONAL ===
        if adapted_cv.get('summary'):
            self._add_section_heading(doc, 'PROFESSIONAL SUMMARY')
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(adapted_cv['summary'])
            summary_run.font.size = Pt(11)
            summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()  # Espacio
        
        # === EXPERIENCIA PROFESIONAL ===
        if adapted_cv.get('experience'):
            self._add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
            
            for i, exp in enumerate(adapted_cv['experience']):
                # Título del puesto y empresa
                job_para = doc.add_paragraph()
                
                # Título del puesto en negrita
                if exp.get('title'):
                    title_run = job_para.add_run(exp['title'])
                    title_run.font.bold = True
                    title_run.font.size = Pt(12)
                    title_run.font.color.rgb = RGBColor(31, 78, 121)
                
                # Empresa
                if exp.get('company'):
                    company_run = job_para.add_run(f" | {exp['company']}")
                    company_run.font.size = Pt(11)
                    company_run.font.color.rgb = RGBColor(89, 89, 89)
                
                # Período en cursiva
                if exp.get('period'):
                    period_para = doc.add_paragraph()
                    period_run = period_para.add_run(exp['period'])
                    period_run.font.italic = True
                    period_run.font.size = Pt(10)
                    period_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Descripción con viñetas mejoradas
                if exp.get('description'):
                    for desc_item in exp['description'][:4]:  # Máximo 4 puntos
                        if isinstance(desc_item, str) and desc_item.strip():
                            bullet_para = doc.add_paragraph()
                            bullet_para.style = 'List Bullet'
                            bullet_run = bullet_para.add_run(f"  {desc_item}")
                            bullet_run.font.size = Pt(10)
                
                # Espacio entre trabajos
                if i < len(adapted_cv['experience']) - 1:
                    doc.add_paragraph()
            
            doc.add_paragraph()  # Espacio después de la sección
        
        # === COMPETENCIAS TÉCNICAS ===
        skills = adapted_cv.get('skills', {})
        if skills.get('technical') or skills.get('other'):
            self._add_section_heading(doc, 'SKILLS')
            
            if skills.get('technical'):
                # Categorizar habilidades técnicas
                tech_categories = self._categorize_technical_skills(skills['technical'])
                
                for category, skill_list in tech_categories.items():
                    if skill_list:
                        cat_para = doc.add_paragraph()
                        cat_run = cat_para.add_run(f"{category}: ")
                        cat_run.font.bold = True
                        cat_run.font.size = Pt(10)
                        
                        skills_run = cat_para.add_run(', '.join(skill_list))
                        skills_run.font.size = Pt(10)
            
            # Competencias adicionales
            if skills.get('other'):
                other_para = doc.add_paragraph()
                other_label = other_para.add_run("Otras competencias: ")
                other_label.font.bold = True
                other_label.font.size = Pt(10)
                
                other_skills = other_para.add_run(', '.join(skills['other']))
                other_skills.font.size = Pt(10)
            
            doc.add_paragraph()  # Espacio
        
        # === EDUCACIÓN ===
        if adapted_cv.get('education'):
            self._add_section_heading(doc, 'EDUCATION')
            
            for edu in adapted_cv['education']:
                if edu.get('degree'):
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"• {edu['degree']}")
                    edu_run.font.size = Pt(10)
                    
                    if edu.get('institution'):
                        inst_run = edu_para.add_run(f" - {edu['institution']}")
                        inst_run.font.italic = True
                        inst_run.font.size = Pt(10)
                        inst_run.font.color.rgb = RGBColor(89, 89, 89)
            
            doc.add_paragraph()
        
        # Guardar documento
        doc.save(filename)
        self.logger.info(f"CV profesional generado: {filename}")
        return filename
    
    def _add_section_heading(self, doc, title: str):
        """Añade un encabezado de sección con formato profesional"""
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(title)
        heading_run.font.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        # Línea inferior del encabezado
        line_para = doc.add_paragraph()
        line_run = line_para.add_run('─' * 60)
        line_run.font.color.rgb = RGBColor(200, 200, 200)
        line_run.font.size = Pt(8)
    
    def _add_horizontal_line(self, doc):
        """Añade una línea horizontal decorativa"""
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run('─' * 40)
        line_run.font.color.rgb = RGBColor(31, 78, 121)
        line_run.font.size = Pt(10)
    
    def _categorize_technical_skills(self, technical_skills: List[str]) -> Dict[str, List[str]]:
        """Categoriza las habilidades técnicas para mejor presentación"""
        categories = {
            'Lenguajes de Programación': [],
            'Bases de Datos': [],
            'Herramientas de Análisis': [],
            'Tecnologías Cloud': [],
            'Frameworks y Librerías': [],
            'Otras Tecnologías': []
        }
        
        # Mapeo de tecnologías a categorías
        tech_mapping = {
            'python': 'Lenguajes de Programación',
            'sql': 'Bases de Datos',
            'javascript': 'Lenguajes de Programación',
            'r': 'Lenguajes de Programación',
            'mysql': 'Bases de Datos',
            'postgresql': 'Bases de Datos',
            'mongodb': 'Bases de Datos',
            'oracle': 'Bases de Datos',
            'power bi': 'Herramientas de Análisis',
            'tableau': 'Herramientas de Análisis',
            'excel': 'Herramientas de Análisis',
            'aws': 'Tecnologías Cloud',
            'azure': 'Tecnologías Cloud',
            'gcp': 'Tecnologías Cloud',
            'tensorflow': 'Frameworks y Librerías',
            'pytorch': 'Frameworks y Librerías',
            'pandas': 'Frameworks y Librerías',
            'numpy': 'Frameworks y Librerías',
            'matplotlib': 'Frameworks y Librerías',
            'scikit-learn': 'Frameworks y Librerías',
            'keras': 'Frameworks y Librerías',
            'git': 'Otras Tecnologías',
            'docker': 'Otras Tecnologías',
            'kubernetes': 'Otras Tecnologías'
        }
        
        for skill in technical_skills:
            skill_lower = skill.lower().strip()
            category = 'Otras Tecnologías'  # Por defecto
            
            # Buscar la categoría apropiada
            for tech, cat in tech_mapping.items():
                if tech in skill_lower:
                    category = cat
                    break
            
            categories[category].append(skill)
        
        # Filtrar categorías vacías
        return {k: v for k, v in categories.items() if v}
    
    # Métodos auxiliares simplificados
    def _adapt_personal_info(self, personal_info: Dict) -> Dict:
        """Adapta la información personal manteniendo los datos esenciales"""
        return personal_info
    
    def _generate_targeted_summary(self, cv: Dict, job_offer: Dict, matching_results: Dict) -> str:
        """Genera un resumen dirigido a la oferta específica usando LM Studio"""
        try:
            # Intentar usar LM Studio primero
            if lmstudio_client and lmstudio_client.available:
                summary = self._generate_summary_with_lm(cv, job_offer, matching_results)
                if summary:
                    self.logger.info("Resumen generado con LM Studio")
                    return summary
            
            # Fallback a lógica basada en reglas
            self.logger.info("Usando generación de resumen basada en reglas")
            return self._generate_summary_fallback(cv, job_offer, matching_results)
            
        except Exception as e:
            self.logger.error(f"Error generando resumen: {e}")
            return "Profesional motivado con experiencia demostrada en múltiples áreas."
    
    def _generate_summary_with_lm(self, cv: Dict, job_offer: Dict, matching_results: Dict) -> str:
        """Genera resumen usando LM Studio"""
        try:
            # Preparar contexto para LM Studio
            experience = cv.get('experience', [])
            skills_data = cv.get('skills', {})
            
            if isinstance(skills_data, dict):
                skills = skills_data.get('technical', [])
            elif isinstance(skills_data, list):
                skills = skills_data
            else:
                skills = []
            
            # Construir prompt
            system_message = """Eres un experto en redacción de CVs profesionales. Tu tarea es crear un resumen profesional personalizado de 3-4 líneas que:
1. Destaque la experiencia más relevante para la oferta
2. Mencione las habilidades técnicas que coinciden con los requisitos
3. Muestre entusiasmo por la posición
4. Sea conciso y profesional
5. Esté en español"""
            
            exp_summary = ""
            if experience:
                exp_summary = "Experiencia:\n"
                for exp in experience[:3]:
                    exp_summary += f"- {exp.get('title', 'N/A')} en {exp.get('company', 'N/A')}\n"
            
            skills_summary = ", ".join(skills[:10]) if skills else "No especificado"
            
            matches_summary = ""
            if matching_results.get('matches'):
                matches_summary = "Coincidencias clave:\n"
                for match in matching_results['matches'][:5]:
                    matches_summary += f"- {match.get('requirement', 'N/A')}\n"
            
            prompt = f"""Genera un resumen profesional de CV para esta persona que aplica a: {job_offer.get('title', 'la posición')}

{exp_summary}

Habilidades: {skills_summary}

{matches_summary}

Requisitos de la oferta:
{', '.join(job_offer.get('must_have', [])[:5])}

Genera SOLO el resumen profesional (3-4 líneas), sin títulos ni explicaciones adicionales:"""
            
            response = lmstudio_client.chat_completion(prompt, system_message, temperature=0.7)
            
            if response:
                # Limpiar respuesta
                summary = response.strip()
                # Limitar longitud
                if len(summary) > 500:
                    summary = summary[:497] + "..."
                return summary
            
            return None
            
        except Exception as e:
            self.logger.error(f"Error generando resumen con LM Studio: {e}")
            return None
    
    def _generate_summary_fallback(self, cv: Dict, job_offer: Dict, matching_results: Dict) -> str:
        """Genera resumen usando lógica basada en reglas (fallback)"""
        experience = cv.get('experience', [])
        skills_data = cv.get('skills', {})
        
        if isinstance(skills_data, dict):
            skills = skills_data.get('technical', [])
        elif isinstance(skills_data, list):
            skills = skills_data
        else:
            skills = []
        
        matches = matching_results.get('matches', [])
        years = self._calculate_experience_years(experience)
        
        summary_parts = []
        
        if years > 0:
            summary_parts.append(f"Profesional con {years} años de experiencia")
        
        main_area = self._identify_main_area(experience)
        if main_area:
            summary_parts.append(f"en {main_area}")
        
        relevant_skills = []
        for skill in skills[:3]:
            if any(skill.lower() in match.get('evidence', '').lower() for match in matches):
                relevant_skills.append(skill)
        
        if relevant_skills:
            summary_parts.append(f"con experiencia en {', '.join(relevant_skills)}")
        
        summary = '. '.join(summary_parts) + '.'
        
        if job_offer.get('title'):
            summary += f" Interesado en contribuir como {job_offer['title']}"
            coverage = matching_results.get('coverage_percentage', 0) / 100
            if coverage > 0.6:
                summary += " aportando experiencia relevante."
            else:
                summary += " y desarrollar nuevas competencias."
        
        return summary
    
    def _calculate_experience_years(self, experience: List[Dict]) -> int:
        """Calcula años totales de experiencia aproximados"""
        return min(len(experience) * 2, 15)  # Estimación simple
    
    def _identify_main_area(self, experience: List[Dict]) -> str:
        """Identifica el área principal de experiencia"""
        if not experience:
            return "tecnología y análisis de datos"
        
        # Buscar patrones comunes en títulos
        titles = [exp.get('title', '').lower() for exp in experience[:3]]
        title_text = ' '.join(titles)
        
        if any(word in title_text for word in ['data', 'datos', 'analista', 'scientist']):
            return "ciencia de datos y análisis"
        elif any(word in title_text for word in ['developer', 'programador', 'software']):
            return "desarrollo de software"
        elif any(word in title_text for word in ['machine learning', 'ai', 'artificial']):
            return "inteligencia artificial"
        else:
            return "tecnología y análisis"
    
    def _prioritize_experience(self, experience: List[Dict], matching_results: Dict) -> List[Dict]:
        """Prioriza las experiencias más relevantes"""
        if not experience:
            return []
        
        # Calcular relevancia para cada experiencia
        matches = matching_results.get('matches', [])
        
        for exp in experience:
            exp['relevance_score'] = 0
            exp_text = json.dumps(exp).lower()
            
            for match in matches:
                evidence = match.get('evidence', '').lower()
                if evidence in exp_text:
                    exp['relevance_score'] += match.get('similarity', 0)
        
        # Ordenar por relevancia
        sorted_exp = sorted(experience, key=lambda x: x.get('relevance_score', 0), reverse=True)
        return sorted_exp[:5]  # Top 5 experiencias
    
    def _highlight_relevant_skills(self, skills: Dict, matching_results: Dict) -> Dict:
        """Destaca las habilidades más relevantes"""
        if not skills:
            return {}
        
        matches = matching_results.get('matches', [])
        matched_skills = set()
        
        # Manejar diferentes formatos de skills
        if isinstance(skills, dict):
            # Identificar skills que aparecen en matches
            for match in matches:
                evidence = match.get('evidence', '').lower()
                for skill_category_key, skill_category in skills.items():
                    if isinstance(skill_category, list):
                        for skill in skill_category:
                            if skill.lower() in evidence:
                                matched_skills.add(skill)
                    elif isinstance(skill_category, str):
                        if skill_category.lower() in evidence:
                            matched_skills.add(skill_category)
            
            # Reorganizar skills destacando los relevantes
            technical_skills = skills.get('technical', [])
            other_skills = skills.get('other', [])
            
            return {
                'technical': technical_skills,
                'other': other_skills,
                'highlighted': list(matched_skills)
            }
        elif isinstance(skills, list):
            # Si skills es una lista, convertir a formato dict
            for match in matches:
                evidence = match.get('evidence', '').lower()
                for skill in skills:
                    if skill.lower() in evidence:
                        matched_skills.add(skill)
            
            return {
                'technical': skills,
                'other': [],
                'highlighted': list(matched_skills)
            }
        else:
            return {'technical': [], 'other': [], 'highlighted': []}
    
    def _highlight_relevant_projects(self, projects: List[Dict], matching_results: Dict) -> List[Dict]:
        """Destaca proyectos relevantes"""
        if not projects:
            return []
        
        matches = matching_results.get('matches', [])
        
        for project in projects:
            project['relevance_score'] = 0
            project_text = json.dumps(project).lower()
            
            for match in matches:
                evidence = match.get('evidence', '').lower()
                if evidence in project_text:
                    project['relevance_score'] += match.get('similarity', 0)
        
        # Ordenar por relevancia
        return sorted(projects, key=lambda x: x.get('relevance_score', 0), reverse=True)
    
    def _generate_adaptation_notes(self, matching_results: Dict) -> Dict:
        """Genera notas sobre la adaptación realizada"""
        matches = matching_results.get('matches', [])
        gaps = matching_results.get('missing_requirements', [])
        coverage = matching_results.get('coverage_percentage', 0) / 100
        
        return {
            'coverage_percentage': round(coverage * 100, 1),
            'matched_requirements': len(matches),
            'gap_requirements': len(gaps),
            'adaptation_strategy': self._describe_adaptation_strategy(coverage),
            'recommendations': self._generate_improvement_recommendations(gaps)
        }
    
    def _describe_adaptation_strategy(self, coverage: float) -> str:
        """Describe la estrategia de adaptación utilizada"""
        if coverage >= 0.8:
            return "CV altamente alineado - contenido reorganizado para destacar experiencias relevantes"
        elif coverage >= 0.6:
            return "CV bien alineado - priorizadas experiencias coincidentes"
        elif coverage >= 0.4:
            return "CV parcialmente alineado - destacadas las coincidencias existentes"
        else:
            return "CV con gaps significativos - organizado para mostrar potencial transferible"
    
    def _generate_improvement_recommendations(self, gaps: List[str]) -> List[str]:
        """Genera recomendaciones de mejora para gaps"""
        recommendations = []
        
        if not gaps:
            return ["¡Excelente! Tu perfil cubre todos los requisitos principales."]
        
        # Tomar los gaps más importantes
        important_gaps = gaps[:3]  # Top 3 gaps
        
        if important_gaps:
            gap_text = ', '.join(str(gap) for gap in important_gaps)
            recommendations.append(f"Considera desarrollar: {gap_text}")
        
        if len(gaps) > 3:
            recommendations.append("Enfócate en los requisitos más importantes para maximizar el impacto")
        
        return recommendations