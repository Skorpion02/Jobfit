import gradio as gr
import os
import json
from typing import Dict, Optional, Tuple, List
import tempfile
from pathlib import Path

# Importar nuestros módulos
from src.scraper.job_scraper import JobScraper
from src.auditor.realism_scorer import RealismScorer
from src.extractor.job_parser import JobParser
from src.extractor.cv_parser import CVParser
from src.matcher.semantic_matcher import SemanticMatcher
from src.generator.cv_adapter import CVAdapter
from src.generator.ats_optimizer import ATS_Optimizer # <- Importar el nuevo optimizador
from src.generator.cv_full_analyzer import CVFullAnalyzer  # <- Análisis Pro ATS
import uuid # <- Importar para nombres de archivo únicos

class JobFitApp:
    def __init__(self):
        self.scraper = JobScraper()
        self.scorer = RealismScorer()
        self.job_parser = JobParser()
        self.cv_parser = CVParser()
        self.matcher = SemanticMatcher()
        self.adapter = CVAdapter()
        self.optimizer = ATS_Optimizer() # <- Instanciar el nuevo optimizador
        self.full_analyzer = CVFullAnalyzer()  # <- Análisis Pro ATS
        
        # Estado de la aplicación
        self.current_job_data = None
        self.current_cv_data = None
        self.current_matching = None
    
    def _clear_cache(self):
        """Limpia el caché de datos previos"""
        self.current_job_data = None
        self.current_cv_data = None
        self.current_matching = None

    # ─────────────────────────────────────────────
    #  ANÁLISIS COMPLETO ATS (Propuesta A → F)
    # ─────────────────────────────────────────────

    def analyze_full_cv(
        self,
        cv_file,
        job_url: str,
        job_text_manual: str,
        idioma: str,
        longitud: str,
        pais: str,
        rol_objetivo: str,
        nivel: str,
    ):
        """
        Ejecuta el análisis completo ATS y devuelve los 6 entregables
        formateados para la interfaz Gradio.

        Returns (tuple of 8):
            tab_a_md, tab_b_md, tab_c_md, tab_d_text, tab_e_md, tab_f_md,
            download_cv_txt, status_md
        """
        try:
            # 1. Obtener texto de la oferta
            if job_url.strip():
                job_text = self.scraper.scrape_job_offer(job_url.strip())
                if not job_text:
                    err = "❌ No se pudo extraer el texto de la URL proporcionada."
                    return err, "", "", "", "", "", None, None, err
            elif job_text_manual.strip():
                job_text = job_text_manual.strip()
            else:
                err = "❌ Proporciona una URL o pega el texto de la oferta."
                return err, "", "", "", "", "", None, None, err

            # 2. Parsear oferta
            job_data = self.job_parser.extract_job_data(job_text)

            # 3. Parsear CV
            if not cv_file:
                err = "❌ Sube un archivo de CV (PDF, DOCX o TXT)."
                return err, "", "", "", "", "", None, None, err

            cv_file_path = cv_file if isinstance(cv_file, str) else cv_file.name
            file_ext = os.path.splitext(cv_file_path)[1][1:].lower()
            cv_data = self.cv_parser.parse_cv(cv_file_path, file_ext)
            if cv_data is None:
                cv_data = {}

            # 4. Ejecutar análisis completo
            # logros y stack se auto-extraen del CV por el LLM en CVFullAnalyzer
            results = self.full_analyzer.analyze(
                cv_data=cv_data,
                job_data=job_data,
                idioma=idioma or "ES",
                longitud=longitud or "2 páginas",
                pais=pais or "",
                rol_objetivo=rol_objetivo or "",
                nivel=nivel or "",
                logros="",
                stack="",
            )

            llm_note = (
                "\n> 🤖 **LM Studio activo** — análisis generado con IA\n"
                if results.get("llm_used")
                else "\n> ⚠️ **LM Studio no disponible** — resultados aproximados (modo básico)\n"
            )

            # ── Formatear A) Diagnóstico ──────────
            a = results["A_diagnosis"]
            score = a.get("score", "N/D")
            score_emoji = self._get_score_color(score) if isinstance(score, int) else "⚪"
            fortalezas_md = "\n".join(
                f"- **{f['fortaleza']}** → *\"{f.get('cita_cv', '')}\"*"
                for f in a.get("fortalezas", [])
            )
            gaps_md = "\n".join(
                f"- ({g.get('impacto','?').upper()}) {g['gap']}"
                for g in a.get("gaps", [])
            )
            tab_a = f"""{llm_note}
## A) Diagnóstico de Encaje

### Resumen de la oferta
{a.get('resumen_oferta', 'N/D')}

### Score de alineación: {score_emoji} **{score}/100**
{a.get('razon_score', '')}

### Top 5 Fortalezas
{fortalezas_md}

### Top 5 Gaps / Riesgos
{gaps_md}
"""

            # ── Formatear B) Keywords ─────────────
            b_keywords = results["B_keywords"].get("keywords", [])
            rows = []
            for kw in b_keywords:
                estado_icon = {"presente": "✅", "debil": "⚠️", "ausente": "❌"}.get(
                    kw.get("estado", ""), "❓"
                )
                rows.append(
                    f"| {estado_icon} {kw.get('keyword','')} "
                    f"| {kw.get('categoria','')} "
                    f"| {kw.get('estado','')} "
                    f"| {kw.get('ubicacion_cv') or '—'} "
                    f"| {kw.get('sugerencia','—')} |"
                )
            kw_table = "| Estado | Keyword | Categoría | Estado | Ubicación en CV | Sugerencia |\n"
            kw_table += "|--------|---------|-----------|--------|-----------------|------------|\n"
            kw_table = "| Keyword | Categoría | Estado | Ubicación en CV | Sugerencia |\n"
            kw_table += "|---------|-----------|--------|-----------------|------------|\n"
            kw_table += "\n".join(rows)
            tab_b = f"""{llm_note}
## B) Keywords ATS\n\n{kw_table}\n"""

            # ── Formatear C) Plan de cambios ──────
            cambios = results["C_changes"].get("cambios", [])
            alto = [c for c in cambios if c.get("prioridad") == "alto"]
            medio = [c for c in cambios if c.get("prioridad") == "medio"]
            opcional = [c for c in cambios if c.get("prioridad") == "opcional"]

            def _fmt_cambios(lista):
                return "\n".join(
                    f"- **[{c.get('seccion','?')}]** {c.get('que_cambiar','')}\n"
                    f"  *Ejemplo:* {c.get('ejemplo','')}\n"
                    f"  *Impacto:* {c.get('impacto', '')}"
                    for c in lista
                )

            tab_c = f"""{llm_note}
## C) Plan de Cambios Priorizado

### 🔴 Alto Impacto
{_fmt_cambios(alto) or 'Ninguno'}

### 🟡 Medio Impacto
{_fmt_cambios(medio) or 'Ninguno'}

### 🟢 Opcional
{_fmt_cambios(opcional) or 'Ninguno'}
"""

            # ── D) CV reescrito (texto plano) ─────
            tab_d = results["D_cv_rewritten"]

            # Crear archivos para descarga
            temp_dir = Path(tempfile.gettempdir()) / "jobfit_exports"
            temp_dir.mkdir(exist_ok=True)
            uid = uuid.uuid4()
            txt_path = temp_dir / f"cv_ats_pro_{uid}.txt"
            with open(txt_path, "w", encoding="utf-8") as fh:
                fh.write(tab_d)
            docx_path = temp_dir / f"cv_ats_pro_{uid}.docx"
            self._plain_text_to_docx(tab_d, str(docx_path))

            # ── Formatear E) Variantes ────────────
            e = results["E_variants"]
            tab_e = f"""{llm_note}
## E) Variante 1 — ATS-first
*(Alta densidad de keywords, natural)*

### Resumen
{e.get('ats_first', {}).get('resumen', 'N/D')}

### Skills
{e.get('ats_first', {}).get('skills', 'N/D')}

---

## E) Variante 2 — Recruiter-first
*(Narrativa diferenciadora, más humana)*

### Resumen
{e.get('recruiter_first', {}).get('resumen', 'N/D')}

### Skills
{e.get('recruiter_first', {}).get('skills', 'N/D')}
"""

            # ── Formatear F) Checklist ────────────
            checklist_items = results["F_checklist"].get("checklist", [])
            checklist_md = "\n".join(
                f"- {'✅' if c.get('estado') == 'ok' else '⚠️' if c.get('estado') == 'advertencia' else '❌'} "
                f"**{c.get('punto', '')}**: {c.get('detalle', '')}"
                for c in checklist_items
            )
            ok_count = sum(1 for c in checklist_items if c.get("estado") == "ok")
            total = len(checklist_items)
            tab_f = f"""{llm_note}
## F) Checklist ATS Final

**Estado general: {ok_count}/{total} puntos OK**

{checklist_md}
"""

            status = f"✅ Análisis completo generado {'con LM Studio' if results.get('llm_used') else '(modo básico — activa LM Studio para mejor calidad)'}"
            return tab_a, tab_b, tab_c, tab_d, tab_e, tab_f, str(txt_path), str(docx_path), status

        except Exception as exc:
            import traceback
            err = f"❌ Error en el análisis: {exc}\n\n```\n{traceback.format_exc()}\n```"
            return err, "", "", "", "", "", None, None, err

    def _plain_text_to_docx(self, text: str, filename: str) -> str:
        """
        Convierte el CV en texto plano (entregable D) a un DOCX limpio y
        compatible con ATS.  Detecta:
        - Líneas de separación (===, ---) → línea decorativa
        - Encabezados de sección EN MAYÚSCULAS → Heading 2 azul
        - Bullets (•, -, *) → List Bullet
        - Resto → párrafo normal
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Márgenes compactos
        for section in doc.sections:
            from docx.shared import Inches
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Estilo base de párrafo
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        BLUE = RGBColor(31, 78, 121)
        GREY = RGBColor(89, 89, 89)

        for line in text.splitlines():
            stripped = line.strip()

            # Separadores visuales → línea gris fina
            if set(stripped) <= {"=", "-", "─"} and len(stripped) > 3:
                p = doc.add_paragraph()
                run = p.add_run("─" * 55)
                run.font.color.rgb = RGBColor(200, 200, 200)
                run.font.size = Pt(7)
                continue

            # Encabezado de sección (toda en mayúsculas, sin bullets)
            if (
                stripped
                and stripped == stripped.upper()
                and len(stripped) > 2
                and not stripped.startswith(("•", "-", "*", "|"))
            ):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = BLUE
                continue

            # Bullet
            if stripped.startswith(("•", "- ", "* ")):
                content = stripped.lstrip("•-* ").strip()
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(content)
                run.font.size = Pt(10.5)
                continue

            # Línea vacía → espacio
            if not stripped:
                doc.add_paragraph()
                continue

            # Línea de contacto (contiene | entre datos)
            if "|" in stripped and len(stripped.split("|")) >= 2:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.color.rgb = GREY
                continue

            # Párrafo normal
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(10.5)

        doc.save(filename)
        return filename

    def get_diagnostics(self):
        """Obtiene información de diagnóstico del sistema"""
        import platform
        from pathlib import Path
        
        diagnostics = []
        
        # Sistema
        diagnostics.append("## 🖥️ Sistema")
        diagnostics.append(f"- **OS:** {platform.system()} {platform.release()}")
        diagnostics.append(f"- **Python:** {platform.python_version()}")
        diagnostics.append("")
        
        # LM Studio
        diagnostics.append("## 🤖 LM Studio")
        try:
            from src.llm.lmstudio_client import lmstudio_client
            from config.settings import settings
            
            diagnostics.append(f"- **URL configurada:** `{settings.lmstudio_base_url}`")
            diagnostics.append(f"- **Modelo configurado:** `{settings.lmstudio_model}`")
            diagnostics.append(f"- **Habilitado:** {'✅ Sí' if settings.use_lmstudio else '❌ No'}")
            diagnostics.append(f"- **Estado:** {'✅ Conectado' if lmstudio_client.available else '❌ No disponible'}")
            
            if lmstudio_client.available:
                diagnostics.append(f"- **Modelo activo:** `{lmstudio_client.model}`")
                models = lmstudio_client.get_available_models()
                if models:
                    diagnostics.append(f"- **Modelos disponibles:** {len(models)}")
                    for i, model in enumerate(models[:5], 1):
                        diagnostics.append(f"  {i}. `{model}`")
                    if len(models) > 5:
                        diagnostics.append(f"  ... y {len(models) - 5} más")
            else:
                diagnostics.append("")
                diagnostics.append("### ⚠️ LM Studio no disponible")
                diagnostics.append("**Posibles causas:**")
                diagnostics.append("1. LM Studio no está instalado")
                diagnostics.append("2. LM Studio no está abierto")
                diagnostics.append("3. El servidor local no está iniciado")
                diagnostics.append("4. Puerto incorrecto o bloqueado")
                diagnostics.append("")
                diagnostics.append("**Solución:**")
                diagnostics.append("1. Descarga LM Studio desde https://lmstudio.ai")
                diagnostics.append("2. Abre LM Studio")
                diagnostics.append("3. Ve a la pestaña 'Local Server'")
                diagnostics.append("4. Carga un modelo y haz click en 'Start Server'")
        except Exception as e:
            diagnostics.append(f"- **Error:** {str(e)}")
        
        diagnostics.append("")
        
        # Logs
        diagnostics.append("## 📋 Logs")
        log_path = Path(__file__).parent.parent / "logs" / "jobfit.log"
        if log_path.exists():
            size_kb = log_path.stat().st_size / 1024
            diagnostics.append(f"- **Archivo:** `{log_path.name}`")
            diagnostics.append(f"- **Tamaño:** {size_kb:.2f} KB")
            diagnostics.append(f"- **Ruta completa:** `{log_path}`")
            diagnostics.append("")
            diagnostics.append("**💡 Para ver logs en tiempo real, ejecuta:**")
            diagnostics.append("```bash")
            diagnostics.append("view_logs.bat")
            diagnostics.append("```")
        else:
            diagnostics.append("- **Estado:** No hay logs todavía")
        
        return "\n".join(diagnostics)
    
    def get_recent_logs(self, num_lines: int = 100):
        """Obtiene las últimas líneas del log"""
        from pathlib import Path
        
        log_path = Path(__file__).parent.parent / "logs" / "jobfit.log"
        
        if not log_path.exists():
            return "📝 No hay logs disponibles todavía.\n\nLos logs aparecerán cuando uses la aplicación."
        
        try:
            # Intentar con UTF-8 primero, luego con otros encodings
            encodings = ['utf-8', 'cp1252', 'latin-1']
            lines = None
            
            for encoding in encodings:
                try:
                    with open(log_path, 'r', encoding=encoding, errors='ignore') as f:
                        lines = f.readlines()
                        break
                except:
                    continue
            
            if lines:
                recent_lines = lines[-num_lines:]
                return "".join(recent_lines)
            else:
                return "❌ No se pudieron leer los logs con ningún encoding compatible"
                
        except Exception as e:
            return f"❌ Error al leer logs: {e}"

    def _get_lmstudio_status(self):
        """Obtiene el estado de LM Studio para mostrar en la interfaz"""
        try:
            from src.llm.lmstudio_client import lmstudio_client
            from config.settings import settings
            
            if settings.use_lmstudio and lmstudio_client.available:
                models = lmstudio_client.get_available_models()
                current_model = lmstudio_client.model
                return f"""
> 🤖 **LM Studio activo** - Modelo: `{current_model}` 
> 📋 Modelos disponibles: {', '.join(models[:3])}{'...' if len(models) > 3 else ''}
"""
            elif settings.use_lmstudio:
                return """
> ⚠️ **LM Studio configurado pero no disponible** - Usando extracción basada en reglas
> 💡 Inicia LM Studio para mejorar la precisión del análisis
"""
            else:
                return """
> 📝 **Modo básico** - Usando extracción basada en reglas
"""
        except:
            return """
> 📝 **Modo básico** - Usando extracción basada en reglas
"""

    def audit_job_offer(self, url: str, manual_text: str) -> Tuple[str, str, str]:
        """Audita una oferta de trabajo"""
        try:
            # Limpiar caché previo
            self._clear_cache()
            
            # Obtener texto de la oferta
            if url.strip():
                job_text = self.scraper.scrape_job_offer(url)
                if not job_text:
                    return "❌ Error: No se pudo extraer texto de la URL", "", ""
            elif manual_text.strip():
                job_text = manual_text
            else:
                return "❌ Error: Proporciona una URL o pega el texto de la oferta", "", ""
            
            # Extraer datos estructurados
            self.current_job_data = self.job_parser.extract_job_data(job_text)
            
            # Calcular score de realismo
            audit_result = self.scorer.calculate_realism_score(self.current_job_data)
            
            # Formatear resultados
            score = audit_result['realism_score']
            score_color = self._get_score_color(score)
            
            # Resumen del audit
            audit_summary = f"""
## 📊 Score de Realismo: {score_color} {score}/100

**{audit_result['reasoning']}**

### 🚨 Señales Detectadas:
"""
            
            for signal in audit_result['signals']:
                emoji = "⚠️" if signal['type'] == 'warning' else "❌" if signal['type'] == 'error' else "ℹ️"
                audit_summary += f"\n{emoji} {signal['description']}"
            
            if not audit_result['signals']:
                audit_summary += "\n✅ No se detectaron problemas significativos"
            
            # Categorías detalladas
            categories = audit_result.get('categories', {})
            category_detail = f"""
### 📈 Análisis por Categorías:
- **Años/Seniority**: {categories.get('years_seniority', 'N/A')}/100
- **Stack Tecnológico**: {categories.get('tech_stack', 'N/A')}/100
- **Salario**: {categories.get('salary', 'N/A') or 'No especificado'}/100
- **Consistencia**: {categories.get('consistency', 'N/A')}/100
"""
            
            # JSON estructurado
            job_json = json.dumps(self.current_job_data, indent=2, ensure_ascii=False)
            
            return audit_summary, category_detail, job_json
            
        except Exception as e:
            return f"❌ Error procesando la oferta: {str(e)}", "", ""
    
    def process_cv_and_match(self, cv_file, job_url: str, manual_job_text: str) -> Tuple[str, str, Optional[str], Optional[str]]:
        """Procesa CV, hace matching y genera AMBOS archivos"""
        try:
            # Limpiar caché previo para evitar datos de consultas anteriores
            self._clear_cache()
            
            # Siempre procesar datos de la oferta actual (evitar caché)
            if job_url.strip():
                job_text = self.scraper.scrape_job_offer(job_url)
                if not job_text:
                    return "❌ Error: No se pudo extraer texto de la URL", "", None, None
            elif manual_job_text.strip():
                job_text = manual_job_text
            else:
                return "❌ Error: Proporciona una URL o pega el texto de la oferta", "", None, None
            
            # Extraer datos de la oferta (sin usar caché)
            job_data = self.job_parser.extract_job_data(job_text)
            
            # Procesar CV
            if not cv_file:
                return "❌ Error: Sube un archivo de CV", "", None, None
            
            # En Gradio 4.4.0, cv_file es directamente la ruta del archivo temporal
            cv_file_path = cv_file if isinstance(cv_file, str) else cv_file.name
            
            # Determinar tipo de archivo
            file_extension = os.path.splitext(cv_file_path)[1][1:].lower()
            
            # Parsear CV
            self.current_cv_data = self.cv_parser.parse_cv(cv_file_path, file_extension)

            # Defensive normalization: asegurar tipos esperados para evitar
            # errores 'NoneType' object is not iterable en la UI
            if self.current_cv_data is None:
                self.current_cv_data = {}

            if not isinstance(self.current_cv_data.get('experience'), list):
                self.current_cv_data['experience'] = self.current_cv_data.get('experience') or []
            if not isinstance(self.current_cv_data.get('education'), list):
                self.current_cv_data['education'] = self.current_cv_data.get('education') or []
            if not isinstance(self.current_cv_data.get('projects'), list):
                self.current_cv_data['projects'] = self.current_cv_data.get('projects') or []

            # Normalizar skills a dict con 'technical' y 'other'
            skills_field = self.current_cv_data.get('skills')
            if isinstance(skills_field, list):
                self.current_cv_data['skills'] = {'technical': skills_field, 'other': []}
            elif isinstance(skills_field, dict):
                tech = skills_field.get('technical') or []
                other = skills_field.get('other') or []
                if isinstance(tech, str):
                    tech = [tech]
                if isinstance(other, str):
                    other = [other]
                self.current_cv_data['skills'] = {'technical': tech, 'other': other}
            else:
                self.current_cv_data['skills'] = {'technical': [], 'other': []}

            # Asegurar raw_text
            if not isinstance(self.current_cv_data.get('raw_text'), str):
                self.current_cv_data['raw_text'] = ''
            
            # Realizar matching con información de tipos de requisitos
            # Normalización defensiva para evitar iteración sobre None
            must_have_reqs = job_data.get('must_have') or []
            if not isinstance(must_have_reqs, list):
                must_have_reqs = []
            
            nice_to_have_reqs = job_data.get('nice_to_have') or []
            if not isinstance(nice_to_have_reqs, list):
                nice_to_have_reqs = []
            
            # Crear lista de requisitos con tipos
            requirements_with_types = []
            for req in must_have_reqs:
                requirements_with_types.append({
                    'requirement': req, 
                    'type': 'must_have'
                })
            for req in nice_to_have_reqs:
                requirements_with_types.append({
                    'requirement': req, 
                    'type': 'nice_to_have'
                })
            
            # Para compatibilidad con el matcher actual
            requirements = must_have_reqs + nice_to_have_reqs
            
            # Debug: verificar que tenemos datos
            print(f"DEBUG - Requirements encontrados: {len(requirements)}")
            print(f"DEBUG - CV data keys: {list(self.current_cv_data.keys())}")
            
            cv_sections = {
                'experience': self._safe_join_cv_section(self.current_cv_data.get('experience', [])),
                # Para skills preferimos concatenar la lista 'technical' si existe
                'skills': self._safe_join_cv_section(self.current_cv_data.get('skills', {}).get('technical', [])),
                'projects': self._safe_join_cv_section(self.current_cv_data.get('projects', [])),
                'education': self._safe_join_cv_section(self.current_cv_data.get('education', []))
            }
            
            # Debug: verificar contenido de CV sections
            for section, content in cv_sections.items():
                print(f"DEBUG - {section}: {len(content)} caracteres")
            
            # Si no tenemos requirements, intentar extraer de texto plano
            if not requirements and job_data.get('raw_text'):
                # Extraer keywords básicos como requirements
                job_text_fallback = job_data.get('raw_text', '')
                requirements = self._extract_basic_requirements(job_text_fallback)
                print(f"DEBUG - Requirements extraídos de texto: {len(requirements)}")
            
            # Concatenar todas las secciones del CV en texto
            cv_text = ' '.join([
                cv_sections['experience'],
                cv_sections['skills'], 
                cv_sections['projects'],
                cv_sections['education']
            ])
            
            self.current_matching = self.matcher.match_requirements_to_cv(
                requirements, cv_text)
            
            # Enriquecer resultados con tipos de requisitos
            self._enrich_matching_with_types(
                self.current_matching, requirements_with_types)
            
            print(f"DEBUG - Matching results: {self.current_matching}")
            
            # Formatear resultados del matching
            matching_summary = self._format_matching_results(
                self.current_matching)
            
            # --- GENERACIÓN DUAL DE CVs ---

            # 1. Generar CV adaptado para humanos (DOCX)
            adapted_cv = self.adapter.adapt_cv(
                self.current_cv_data, 
                job_data, 
                self.current_matching
            )
            cv_preview = self._format_cv_preview(adapted_cv)
            
            # Crear directorio temporal si no existe
            temp_dir = Path(tempfile.gettempdir()) / "jobfit_exports"
            temp_dir.mkdir(exist_ok=True)

            # Exportar DOCX a archivo temporal
            temp_docx_path = temp_dir / f"cv_adaptado_{uuid.uuid4()}.docx"
            self.adapter.export_to_docx(adapted_cv, str(temp_docx_path))

            # 2. Generar CV optimizado para ATS (TXT)
            ats_optimized_text = self.optimizer.optimize_cv_for_ats(
                self.current_cv_data, self.current_matching
            )
            temp_txt_path = temp_dir / f"cv_optimizado_ats_{uuid.uuid4()}.txt"
            with open(temp_txt_path, 'w', encoding='utf-8') as f:
                f.write(ats_optimized_text)

            return matching_summary, cv_preview, str(temp_docx_path), str(temp_txt_path)
            
        except Exception as e:
            return f"❌ Error procesando CV: {str(e)}", "", None, None
    
    def _get_score_color(self, score: int) -> str:
        """Devuelve emoji de color según el score"""
        if score >= 85:
            return "🟢"
        elif score >= 70:
            return "🟡"
        elif score >= 50:
            return "🟠"
        else:
            return "🔴"
    
    def _format_matching_results(self, matching: Dict) -> str:
        """Formatea los resultados del matching"""
        matches = matching.get('matches', [])
        gaps = matching.get('missing_requirements', [])
        coverage = matching.get('coverage_percentage', 0) / 100
        
        # Validar que tenemos datos
        if coverage == 0 and not matches and not gaps:
            return """
## ⚠️ Sin Resultados de Matching

No se pudieron procesar los requisitos o el CV. Verifica que:
- La oferta de trabajo contenga requisitos claros
- El CV tenga información relevante

**Sugerencia**: Intenta con una oferta más detallada o un CV más completo.
"""
        
        result = f"""
## 🎯 Análisis de Encaje: {coverage*100:.1f}% de cobertura

### ✅ Requisitos Cubiertos ({len(matches)}):
"""
        
        if matches:
            for match in matches[:10]:  # Mostrar top 10
                # Emojis según similitud
                if match['similarity'] > 0.8:
                    confidence_emoji = "🔥"
                elif match['similarity'] > 0.6:
                    confidence_emoji = "✅"
                else:
                    confidence_emoji = "⚡"
                
                # Emoji según tipo de match
                match_type_emoji = ("🎯" if match.get('match_type') == 'exact'
                                    else "🔍")
                
                # Emoji según tipo de requisito
                req_type = match.get('requirement_type', 'unknown')
                if req_type == 'must_have':
                    type_emoji = "🔴"  # Rojo para obligatorios
                    type_text = "Obligatorio"
                elif req_type == 'nice_to_have':
                    type_emoji = "🟡"  # Amarillo para deseables
                    type_text = "Deseable"
                else:
                    type_emoji = "⚪"
                    type_text = "Otro"
                
                evidence = match.get('evidence',
                                    'Detectado por similitud semántica')
                section = match.get('section', 'N/A')
                similarity = match['similarity']
                
                result += f"""
{confidence_emoji} **{match['requirement']}** {type_emoji} *{type_text}*
   {match_type_emoji} *Evidencia*: {evidence}
   📍 *Sección*: {section} | 💪 *Similitud*: {similarity:.1%}
"""
        else:
            result += "\n*No se encontraron matches con suficiente confianza.*\n"
        
        if gaps:
            result += f"\n\n### ❌ Requisitos Faltantes ({len(gaps)}):\n"
            for gap in gaps[:8]:  # Mostrar top 8 gaps
                if isinstance(gap, dict):
                    req_text = gap.get('requirement', gap)
                    req_type = gap.get('requirement_type', 'unknown')
                    if req_type == 'must_have':
                        type_emoji = "🔴"
                        type_text = "Obligatorio"
                    elif req_type == 'nice_to_have':
                        type_emoji = "🟡"
                        type_text = "Deseable"
                    else:
                        type_emoji = "⚪"
                        type_text = "Otro"
                    result += f"• {req_text} {type_emoji} *{type_text}*\n"
                else:
                    result += f"• {gap}\n"
            
            if len(gaps) > 8:
                result += f"• ... y {len(gaps) - 8} requisitos más\n"
                
            result += ("\n**💡 Recomendación**: Considera fortalecer "
                      "estas áreas para mejorar tu perfil.\n")
        else:
            result += ("\n\n### 🎉 ¡Excelente!\n"
                      "Tu perfil cubre todos los requisitos identificados.\n")
        
        return result
    
    def _extract_basic_requirements(self, job_text: str) -> List[str]:
        """Extrae requisitos básicos del texto de la oferta"""
        import re
        
        # Buscar tecnologías comunes
        tech_patterns = [
            r'\b(python|java|javascript|react|angular|vue|node\.?js)\b',
            r'\b(html|css|sql|mysql|postgresql|mongodb)\b',
            r'\b(docker|kubernetes|aws|azure|git|github)\b',
            r'\b(spring|django|flask|express)\b'
        ]
        
        requirements = []
        job_lower = job_text.lower()
        
        for pattern in tech_patterns:
            matches = re.findall(pattern, job_lower, re.IGNORECASE)
            requirements.extend(matches)
        
        # Buscar frases de requisitos
        requirement_patterns = [
            r'experiencia (?:en|con) ([^.,\n]+)',
            r'conocimientos? (?:en|de) ([^.,\n]+)',
            r'dominio de ([^.,\n]+)',
            r'manejo de ([^.,\n]+)'
        ]
        
        for pattern in requirement_patterns:
            matches = re.findall(pattern, job_lower)
            # Filtrar solo strings para evitar error dict.strip()
            string_matches = [match.strip() for match in matches
                              if isinstance(match, str)]
            requirements.extend(string_matches)
        
        # Limpiar y limitar
        clean_requirements = []
        for req in requirements:
            if len(req) > 2 and len(req) < 50:
                clean_requirements.append(req)
        
        return list(set(clean_requirements))[:10]  # Max 10 requisitos únicos
    
    def _safe_join_cv_section(self, section_data) -> str:
        """Une de forma segura los datos de una sección del CV"""
        if not section_data:
            return ""
        
        texts = []
        for item in section_data:
            if isinstance(item, str):
                texts.append(item)
            elif isinstance(item, dict):
                # Si es un diccionario, extraer valores de texto
                for key, value in item.items():
                    if isinstance(value, str):
                        texts.append(value)
                    elif isinstance(value, list):
                        for subitem in value:
                            if isinstance(subitem, str):
                                texts.append(subitem)
            elif isinstance(item, list):
                # Si es una lista, procesar recursivamente
                for subitem in item:
                    if isinstance(subitem, str):
                        texts.append(subitem)
            else:
                # Para cualquier otro tipo, convertir a string
                texts.append(str(item))
        
        return ' '.join(texts)
    
    def _format_cv_preview(self, adapted_cv: Dict) -> str:
        """Formatea vista previa del CV adaptado"""
        preview = f"""
# 📄 Vista Previa del CV Adaptado

## 👤 {adapted_cv.get('personal_info', {}).get('name', 'Nombre')}

### 📝 Resumen Profesional
{adapted_cv.get('summary', 'No disponible')}

### 💼 Experiencia Profesional (Reordenada por Relevancia)
"""
        
        for i, exp in enumerate(adapted_cv.get('experience', [])[:3]):
            preview += f"""
**{i+1}. {exp.get('title', 'Título')} - {exp.get('company', 'Empresa')}**
*{exp.get('period', 'Periodo')}*
"""
            if exp.get('description'):
                for desc in exp['description'][:2]:
                    preview += f"• {desc}\n"
        
        # Skills reorganizadas
        skills = adapted_cv.get('skills', {})
        if skills.get('technical'):
            preview += f"\n### 🛠️ Habilidades Técnicas Destacadas\n{', '.join(skills['technical'][:8])}\n"
        
        # Notas de adaptación
        notes = adapted_cv.get('adaptation_notes', {})
        if notes:
            preview += f"""
### 📊 Notas de Adaptación
- **Cobertura**: {notes.get('coverage_percentage', 0)}%
- **Estrategia**: {notes.get('adaptation_strategy', 'N/A')}
- **Recomendaciones**: {'; '.join(notes.get('recommendations', [])[:2])}
"""
        
        return preview
    
    def _enrich_matching_with_types(self, matching_results: Dict,
                                    requirements_with_types: List[Dict]):
        """Enriquece los resultados del matching con información de tipos"""
        type_map = {}
        for req_info in requirements_with_types:
            type_map[req_info['requirement']] = req_info['type']
        
        # Agregar tipo a los matches
        for match in matching_results.get('matches', []):
            req_text = match['requirement']
            match['requirement_type'] = type_map.get(req_text, 'unknown')
        
        # Agregar tipo a los missing requirements
        enriched_missing = []
        for req in matching_results.get('missing_requirements', []):
            enriched_missing.append({
                'requirement': req,
                'requirement_type': type_map.get(req, 'unknown')
            })
        matching_results['missing_requirements'] = enriched_missing
    
    def create_interface(self):
        """Crea la interfaz Gradio"""
        
        # Verificar estado de LM Studio
        lmstudio_status = self._get_lmstudio_status()
        
        with gr.Blocks(title="JobFit Agent") as app:
            gr.Markdown(f"""
# 🎯 JobFit Agent
## Auditoría Inteligente de Ofertas y Adaptación de CV

Analiza ofertas de trabajo y adapta tu CV de forma inteligente sin añadir información falsa.

{lmstudio_status}
""")
            
            with gr.Tabs():
                # Pestaña 1: Auditoría de Oferta
                with gr.TabItem("🔍 Auditar Oferta"):
                    gr.Markdown("### Analiza el realismo y extrae información estructurada de una oferta")
                    
                    with gr.Row():
                        with gr.Column():
                            job_url = gr.Textbox(
                                label="🔗 URL de la oferta", 
                                placeholder="https://...",
                                lines=1
                            )
                            job_text_manual = gr.Textbox(
                                label="📝 O pega aquí el texto de la oferta", 
                                placeholder="Descripción completa de la oferta...",
                                lines=8
                            )
                            audit_btn = gr.Button("🚀 Auditar Oferta", variant="primary")
                        
                        with gr.Column():
                            audit_results = gr.Markdown(label="📊 Resultados del Audit")
                            category_breakdown = gr.Markdown(label="📈 Desglose por Categorías")
                    
                    with gr.Row():
                        job_json_output = gr.Code(
                            label="📋 Datos Estructurados (JSON)", 
                            language="json"
                        )
                    
                    audit_btn.click(
                        fn=self.audit_job_offer,
                        inputs=[job_url, job_text_manual],
                        outputs=[audit_results, category_breakdown, job_json_output]
                    )
                    
                # Pestaña 3: Análisis Pro ATS (A → F)
                with gr.TabItem("🎯 Análisis Pro ATS"):
                    gr.Markdown("""
### Análisis Completo ATS
Genera los **6 entregables** del proceso de adaptación profesional:
A) Diagnóstico · B) Keywords · C) Plan de cambios · D) CV reescrito · E) Variantes · F) Checklist
""")
                    with gr.Row():
                        with gr.Column(scale=1):
                            gr.Markdown("#### 📁 Datos de entrada")
                            pro_cv_file = gr.File(
                                label="CV (PDF, DOCX o TXT)",
                                file_types=[".pdf", ".docx", ".txt"],
                            )
                            pro_job_url = gr.Textbox(
                                label="🔗 URL de la oferta (opcional)",
                                placeholder="https://...",
                                lines=1,
                            )
                            pro_job_text = gr.Textbox(
                                label="📝 Texto de la oferta",
                                placeholder="Pega aquí la descripción completa de la oferta...",
                                lines=7,
                            )
                            gr.Markdown("#### ⚙️ Preferencias del candidato")
                            pro_idioma = gr.Dropdown(
                                choices=["ES", "EN", "FR", "DE", "PT"],
                                value="ES",
                                label="Idioma del CV final",
                            )
                            pro_longitud = gr.Dropdown(
                                choices=["1 página", "2 páginas"],
                                value="2 páginas",
                                label="Longitud deseada",
                            )
                            pro_pais = gr.Textbox(
                                label="País / mercado (opcional)",
                                placeholder="Ej: España, Latinoamérica…",
                                lines=1,
                            )
                            pro_rol = gr.Textbox(
                                label="Rol objetivo exacto (si difiere de la oferta)",
                                placeholder="Ej: Data Analyst Senior",
                                lines=1,
                            )
                            pro_nivel = gr.Dropdown(
                                choices=["", "Junior", "Mid", "Senior", "Lead / Principal"],
                                value="",
                                label="Nivel deseado",
                            )
                            gr.Markdown(
                                "> 🤖 Los **logros medibles** y el **stack tecnológico** "
                                "se extraen automáticamente de tu CV con LM Studio."
                            )
                            pro_btn = gr.Button("🚀 Generar Análisis Completo", variant="primary", size="lg")

                        with gr.Column(scale=2):
                            pro_status = gr.Markdown(label="Estado")
                            with gr.Tabs():
                                with gr.TabItem("A) Diagnóstico"):
                                    pro_tab_a = gr.Markdown()
                                with gr.TabItem("B) Keywords ATS"):
                                    pro_tab_b = gr.Markdown()
                                with gr.TabItem("C) Plan de cambios"):
                                    pro_tab_c = gr.Markdown()
                                with gr.TabItem("D) CV reescrito"):
                                    pro_tab_d = gr.Textbox(
                                        label="CV en texto plano (listo para pegar)",
                                        lines=35,
                                    )
                                    with gr.Row():
                                        pro_download = gr.File(
                                            label="⬇️ Descargar TXT",
                                            interactive=False,
                                        )
                                        pro_download_docx = gr.File(
                                            label="⬇️ Descargar DOCX",
                                            interactive=False,
                                        )
                                with gr.TabItem("E) Variantes"):
                                    pro_tab_e = gr.Markdown()
                                with gr.TabItem("F) Checklist"):
                                    pro_tab_f = gr.Markdown()

                    pro_btn.click(
                        fn=self.analyze_full_cv,
                        inputs=[
                            pro_cv_file,
                            pro_job_url,
                            pro_job_text,
                            pro_idioma,
                            pro_longitud,
                            pro_pais,
                            pro_rol,
                            pro_nivel,
                        ],
                        outputs=[
                            pro_tab_a,
                            pro_tab_b,
                            pro_tab_c,
                            pro_tab_d,
                            pro_tab_e,
                            pro_tab_f,
                            pro_download,
                            pro_download_docx,
                            pro_status,
                        ],
                    )

                # Pestaña 4: Diagnóstico
                with gr.TabItem("🔧 Diagnóstico"):
                    gr.Markdown("### Diagnóstico del Sistema y LM Studio")
                    
                    with gr.Row():
                        refresh_diagnostics_btn = gr.Button("🔄 Actualizar Diagnóstico", variant="secondary")
                    
                    diagnostics_output = gr.Markdown(value=self.get_diagnostics())
                    
                    gr.Markdown("### 📋 Logs Recientes")
                    
                    with gr.Row():
                        with gr.Column(scale=4):
                            log_lines_slider = gr.Slider(
                                minimum=50, maximum=500, value=100, step=50,
                                label="Número de líneas a mostrar"
                            )
                        with gr.Column(scale=1):
                            refresh_logs_btn = gr.Button("🔄 Actualizar Logs", variant="secondary")
                    
                    logs_output = gr.Textbox(
                        value=self.get_recent_logs(100),
                        label="Últimas líneas del log",
                        lines=20,
                        max_lines=30
                    )
                    
                    gr.Markdown("""
### 💡 Consejos de Diagnóstico

**Si LM Studio no se conecta:**
1. Verifica que LM Studio esté abierto
2. Ve a la pestaña "Local Server" en LM Studio
3. Asegúrate de que un modelo esté cargado
4. Haz click en "Start Server"
5. Verifica que el puerto sea 1234 (o actualiza `config/settings.py`)

**Para ver logs en tiempo real:**
- Ejecuta `view_logs.bat` en una ventana separada
- Los logs se actualizarán automáticamente

**Niveles de log:**
- 🟢 INFO: Operaciones normales
- 🟡 WARNING: Advertencias no críticas
- 🔴 ERROR: Errores que requieren atención
""")
                    
                    # Eventos de actualización
                    refresh_diagnostics_btn.click(
                        fn=self.get_diagnostics,
                        inputs=[],
                        outputs=[diagnostics_output]
                    )
                    
                    refresh_logs_btn.click(
                        fn=self.get_recent_logs,
                        inputs=[log_lines_slider],
                        outputs=[logs_output]
                    )
                
                # Pestaña 4: Información
                with gr.TabItem("ℹ️ Información"):
                    gr.Markdown("""
### 🤖 Sobre JobFit Agent

Este agente inteligente te ayuda a:

1. **🔍 Auditar ofertas de trabajo**: 
   - Calcula un score de realismo (0-100)
   - Detecta requisitos excesivos o contradictorios
   - Extrae información estructurada

2. **🎯 Adaptar tu CV**: 
   - Reordena experiencias por relevancia
   - Genera resumen dirigido a la oferta
   - **NUNCA inventa información falsa**

3. **📊 Análisis de encaje**:
   - Matching semántico entre requisitos y evidencias
   - Identifica gaps de competencias
   - Proporciona recomendaciones de mejora

### 🛡️ Principios de Veracidad
- ✅ Solo reorganiza información real del CV
- ✅ Destaca experiencias relevantes existentes
- ❌ NO inventa skills ni experiencias
- ❌ NO añade información no verificable

### 🔒 Privacidad
- Los datos se procesan localmente
- No se almacenan CVs ni información personal
- Archivos temporales se eliminan automáticamente
""")
        
        return app

def launch_app():
    """Lanza la aplicación"""
    app = JobFitApp()
    interface = app.create_interface()
    
    # Try multiple ports to avoid conflicts
    ports_to_try = [7860, 7861, 7862, 7863, 7864]
    
    for port in ports_to_try:
        try:
            interface.launch(
                server_name="0.0.0.0",
                server_port=port,
                share=False,
                debug=True,
                theme=gr.themes.Soft()
            )
            break
        except Exception as e:
            if "Cannot find empty port" in str(e) and port != ports_to_try[-1]:
                print(f"Port {port} is busy, trying next port...")
                continue
            else:
                raise e


if __name__ == "__main__":
    launch_app()